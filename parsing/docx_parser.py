from typing import List, Dict
from docx import Document
from parsing.base import BaseParser


class DocxParser(BaseParser):

    def parse_flat(self, filepath: str) -> str:
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)

    def parse_structured(self, filepath: str) -> List[Dict]:
        doc = Document(filepath)
        elements = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                elements.append({
                    "text": t,
                    "style": para.style.name,
                    "is_heading": para.style.name.lower().startswith("heading"),
                })
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    elements.append({"text": row_text, "style": "Table", "is_heading": False})
        return elements
