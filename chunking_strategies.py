"""
Chunking strategies for the RAG benchmark.
Strategies: sliding_window (C, D), structure_aware (E), semantic (F)
"""

import re
from typing import List, Dict
import numpy as np
import tiktoken
from docx import Document
from sentence_transformers import SentenceTransformer

# Shared tokenizer
TOKENIZER = tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str) -> int:
    return len(TOKENIZER.encode(text))


def load_docx_text(filepath: str) -> str:
    """Load plain text from a .docx file."""
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def load_docx_structured(filepath: str) -> List[Dict]:
    """Load document as list of {text, style, is_heading} dicts."""
    doc = Document(filepath)
    elements = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            elements.append({
                "text": t,
                "style": para.style.name,
                "is_heading": para.style.name.lower().startswith("heading"),
            })
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                elements.append({
                    "text": row_text,
                    "style": "Table",
                    "is_heading": False,
                })
    return elements


# ─── Strategy C & D: Sliding Window ──────────────────────────────────────────

def sliding_window_chunks(
    text: str,
    chunk_size: int,
    overlap: int,
    doc_name: str,
) -> List[Dict]:
    """Token-based sliding window chunking."""
    tokens = TOKENIZER.encode(text)
    chunks = []
    start = 0
    idx = 0

    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = TOKENIZER.decode(chunk_tokens)

        chunks.append({
            "chunk_id": f"{doc_name}::sw_{idx}",
            "text": chunk_text,
            "source_doc": doc_name,
            "n_tokens": len(chunk_tokens),
        })

        if end >= len(tokens):
            break

        start = end - overlap
        idx += 1

    return chunks


# ─── Strategy E: Structure-Aware ──────────────────────────────────────────────

def structure_aware_chunks(
    elements: List[Dict],
    max_tokens: int,
    min_tokens: int,
    doc_name: str,
) -> List[Dict]:
    """Group paragraphs by structural boundaries (headings) respecting token limits."""
    chunks = []
    current_texts: List[str] = []
    current_tokens = 0
    idx = 0

    def flush():
        nonlocal idx, current_texts, current_tokens
        if not current_texts:
            return
        text = "\n\n".join(current_texts)
        chunks.append({
            "chunk_id": f"{doc_name}::sa_{idx}",
            "text": text,
            "source_doc": doc_name,
            "n_tokens": current_tokens,
        })
        idx += 1
        current_texts = []
        current_tokens = 0

    for elem in elements:
        elem_tokens = count_tokens(elem["text"])

        # Heading triggers a flush if we have enough content
        if elem["is_heading"] and current_tokens >= min_tokens:
            flush()

        # Would overflow max → flush first
        if current_tokens + elem_tokens > max_tokens and current_tokens >= min_tokens:
            flush()

        current_texts.append(elem["text"])
        current_tokens += elem_tokens

    flush()
    return chunks


# ─── Strategy F: Semantic ─────────────────────────────────────────────────────

_embed_model_cache: Dict[str, SentenceTransformer] = {}


def _get_embed_model(model_name: str) -> SentenceTransformer:
    if model_name not in _embed_model_cache:
        _embed_model_cache[model_name] = SentenceTransformer(model_name)
    return _embed_model_cache[model_name]


def semantic_chunks(
    text: str,
    threshold: float,
    max_tokens: int,
    embed_model_name: str,
    doc_name: str,
) -> List[Dict]:
    """Semantic chunking: split when cosine similarity drops below threshold."""
    # Split into sentences
    sentences = re.split(r"(?<=[.!?])\s+", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return []

    model = _get_embed_model(embed_model_name)
    embeddings = model.encode(sentences, show_progress_bar=False, normalize_embeddings=True)
    # embeddings are L2-normalised → dot product = cosine similarity

    chunks = []
    current_sents: List[str] = [sentences[0]]
    current_tokens = count_tokens(sentences[0])
    idx = 0

    for i in range(1, len(sentences)):
        sim = float(np.dot(embeddings[i], embeddings[i - 1]))
        sent_tokens = count_tokens(sentences[i])

        should_split = (sim < threshold) or (current_tokens + sent_tokens > max_tokens)

        if should_split:
            chunk_text = " ".join(current_sents)
            chunks.append({
                "chunk_id": f"{doc_name}::sem_{idx}",
                "text": chunk_text,
                "source_doc": doc_name,
                "n_tokens": current_tokens,
            })
            idx += 1
            current_sents = [sentences[i]]
            current_tokens = sent_tokens
        else:
            current_sents.append(sentences[i])
            current_tokens += sent_tokens

    # Last chunk
    if current_sents:
        chunk_text = " ".join(current_sents)
        chunks.append({
            "chunk_id": f"{doc_name}::sem_{idx}",
            "text": chunk_text,
            "source_doc": doc_name,
            "n_tokens": current_tokens,
        })

    return chunks
